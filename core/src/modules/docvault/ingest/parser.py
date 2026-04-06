"""DocumentParserOp — parse PDF/DOCX/HTML/Markdown to structured markdown.

Uses pdfplumber for PDF and python-docx for DOCX.
Extracts text, tables, headings, and page numbers.
"""

import logging
import re
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


class DocumentParserOp:
    """Fixed Op: raw file bytes → markdown + metadata."""

    @property
    def name(self) -> str:
        return "document_parser"

    @property
    def input_keys(self) -> tuple[str, ...]:
        return ("file_bytes", "source_type")

    @property
    def output_keys(self) -> tuple[str, ...]:
        return ("raw_content", "doc_metadata")

    async def __call__(self, ctx: dict[str, Any]) -> dict[str, Any]:
        file_bytes: bytes = ctx["file_bytes"]
        source_type: str = ctx["source_type"]

        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "markdown": self._parse_markdown,
            "md": self._parse_markdown,
            "txt": self._parse_plaintext,
        }

        parser = parser_map.get(source_type, self._parse_plaintext)
        raw_content, metadata = parser(file_bytes)

        ctx["raw_content"] = raw_content
        ctx["doc_metadata"] = metadata
        logger.info(
            "DocumentParser: source_type=%s, length=%d chars, pages=%s",
            source_type,
            len(raw_content),
            metadata.get("page_count", "n/a"),
        )
        return ctx

    def _parse_pdf(self, data: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            logger.error("pdfplumber not installed — pip install pdfplumber")
            return "", {"error": "pdfplumber not installed"}

        sections: list[str] = []
        metadata: dict[str, Any] = {"page_count": 0}

        with pdfplumber.open(BytesIO(data)) as pdf:
            metadata["page_count"] = len(pdf.pages)
            if pdf.metadata:
                metadata["author"] = pdf.metadata.get("Author", "")
                metadata["title"] = pdf.metadata.get("Title", "")

            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    sections.append(f"<!-- page {i + 1} -->\n{text}")

                # Extract tables as markdown
                for table in page.extract_tables():
                    if table:
                        md_table = self._table_to_markdown(table)
                        sections.append(md_table)

        return "\n\n".join(sections), metadata

    def _parse_docx(self, data: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
        except ImportError:
            logger.error("python-docx not installed — pip install python-docx")
            return "", {"error": "python-docx not installed"}

        doc = docx.Document(BytesIO(data))
        metadata: dict[str, Any] = {}

        if doc.core_properties:
            metadata["author"] = doc.core_properties.author or ""
            metadata["title"] = doc.core_properties.title or ""

        sections: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                sections.append(f"# {text}")
            elif "heading 2" in style_name:
                sections.append(f"## {text}")
            elif "heading 3" in style_name:
                sections.append(f"### {text}")
            elif "heading" in style_name:
                sections.append(f"#### {text}")
            else:
                sections.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                sections.append(self._table_to_markdown(rows))

        metadata["paragraph_count"] = len(doc.paragraphs)
        return "\n\n".join(sections), metadata

    def _parse_markdown(self, data: bytes) -> tuple[str, dict[str, Any]]:
        """Pass through markdown content."""
        text = data.decode("utf-8", errors="replace")
        heading_count = len(re.findall(r"^#{1,6}\s", text, re.MULTILINE))
        return text, {"heading_count": heading_count, "format": "markdown"}

    def _parse_plaintext(self, data: bytes) -> tuple[str, dict[str, Any]]:
        """Parse plain text."""
        text = data.decode("utf-8", errors="replace")
        return text, {"format": "plaintext"}

    @staticmethod
    def _table_to_markdown(rows: list[list[str | None]]) -> str:
        """Convert a list of rows to markdown table format."""
        if not rows:
            return ""

        # Clean cells
        clean_rows = []
        for row in rows:
            clean_rows.append([(cell or "").replace("|", "\\|").replace("\n", " ") for cell in row])

        # Build markdown
        header = "| " + " | ".join(clean_rows[0]) + " |"
        separator = "| " + " | ".join("---" for _ in clean_rows[0]) + " |"
        body_lines = []
        for row in clean_rows[1:]:
            # Pad row to match header width
            while len(row) < len(clean_rows[0]):
                row.append("")
            body_lines.append("| " + " | ".join(row[: len(clean_rows[0])]) + " |")

        return "\n".join([header, separator, *body_lines])
